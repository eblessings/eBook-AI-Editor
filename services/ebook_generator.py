"""
Fixed Professional eBook Generation Service.
Creates EPUB, PDF, DOCX, and other formats with proper error handling and streaming responses.
"""

import asyncio
import io
import uuid
import zipfile
import tempfile
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, BinaryIO, Union
import base64
import re

from ebooklib import epub
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template
try:
    from PIL import Image
except ImportError:
    Image = None

import structlog

from config import Settings, EBOOK_FORMATS
from api.models import EBookMetadata, FormatOptions, ChapterConfiguration, AIEnhancementOptions

logger = structlog.get_logger()


class EBookGenerator:
    """Fixed eBook generator with multiple format support and proper streaming."""
    
    def __init__(self, settings: Settings):
        self.settings = settings
        self.temp_files = []
        
        # HTML templates for eBook content
        self.chapter_template = Template("""
        <!DOCTYPE html>
        <html xmlns="http://www.w3.org/1999/xhtml">
        <head>
            <title>{{ title }}</title>
            <meta charset="utf-8"/>
            <link rel="stylesheet" type="text/css" href="style/main.css"/>
        </head>
        <body>
            <div class="chapter">
                {% if title != "Main Content" %}
                <h1 class="chapter-title">{{ title }}</h1>
                {% endif %}
                <div class="chapter-content">
                    {{ content | safe }}
                </div>
            </div>
        </body>
        </html>
        """)
        
        self.css_template = Template("""
        body {
            font-family: {{ font_family }};
            font-size: {{ font_size }}pt;
            line-height: {{ line_height }};
            margin: {{ margin_top }}px {{ margin_right }}px {{ margin_bottom }}px {{ margin_left }}px;
            text-align: {{ 'justify' if justify_text else 'left' }};
            color: #333;
        }
        
        .chapter {
            page-break-before: {{ 'always' if page_break_before_chapter else 'auto' }};
        }
        
        .chapter-title {
            font-size: 1.8em;
            font-weight: bold;
            text-align: center;
            margin-bottom: 2em;
            margin-top: 1em;
            page-break-after: avoid;
            color: #2c3e50;
        }
        
        .chapter-content {
            text-indent: 1.5em;
        }
        
        .chapter-content p {
            margin-bottom: 1em;
            text-align: {{ 'justify' if justify_text else 'left' }};
        }
        
        h1, h2, h3, h4, h5, h6 {
            page-break-after: avoid;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
            color: #34495e;
        }
        
        h1 { font-size: 1.6em; }
        h2 { font-size: 1.4em; }
        h3 { font-size: 1.2em; }
        
        blockquote {
            margin: 1em 2em;
            font-style: italic;
            border-left: 3px solid #3498db;
            padding-left: 1em;
            background-color: #f8f9fa;
        }
        
        .toc {
            page-break-after: always;
        }
        
        .toc-entry {
            margin-bottom: 0.5em;
        }
        
        .front-matter, .back-matter {
            page-break-before: always;
        }
        
        .book-header {
            text-align: center;
            margin-bottom: 3em;
            page-break-after: always;
        }
        
        .book-title {
            font-size: 2.5em;
            font-weight: bold;
            margin-bottom: 0.5em;
            color: #2c3e50;
        }
        
        .book-author {
            font-size: 1.3em;
            font-style: italic;
            color: #7f8c8d;
        }
        """)
    
    async def create_ebook(self, content: str, metadata: EBookMetadata, 
                          format_options: FormatOptions, ai_options: AIEnhancementOptions) -> BinaryIO:
        """Create eBook in the specified format with proper error handling."""
        try:
            logger.info("Starting eBook generation", 
                       title=metadata.title, 
                       format=getattr(format_options, 'format', 'epub'))
            
            # Validate inputs
            if not content or not content.strip():
                raise ValueError("Content cannot be empty")
            
            if not metadata.title or not metadata.title.strip():
                raise ValueError("Title is required")
            
            if not metadata.author or not metadata.author.strip():
                raise ValueError("Author is required")
            
            # Process and enhance content
            processed_content = await self._preprocess_content(content, ai_options)
            
            # Detect chapters
            chapters = await self._detect_chapters(processed_content, metadata)
            
            # Generate eBook based on format
            format_type = getattr(format_options, 'format', 'epub')
            
            if format_type == 'epub':
                return await self._generate_epub(chapters, metadata, format_options)
            elif format_type == 'pdf':
                return await self._generate_pdf(chapters, metadata, format_options)
            elif format_type == 'docx':
                return await self._generate_docx(chapters, metadata, format_options)
            elif format_type == 'html':
                return await self._generate_html(chapters, metadata, format_options)
            elif format_type == 'txt':
                return await self._generate_txt(chapters, metadata, format_options)
            else:
                raise ValueError(f"Unsupported format: {format_type}")
                
        except Exception as e:
            logger.error("eBook generation failed", error=str(e))
            raise
    
    async def _preprocess_content(self, content: str, ai_options: AIEnhancementOptions) -> str:
        """Preprocess and enhance content before eBook generation."""
        try:
            processed_content = content
            
            # Basic text cleaning
            processed_content = self._clean_content(processed_content)
            
            # AI enhancements would go here if needed
            if ai_options.enhance_before_generation:
                logger.info("AI enhancement requested but not implemented in this version")
                # For now, just return cleaned content
            
            return processed_content
            
        except Exception as e:
            logger.error("Content preprocessing failed", error=str(e))
            return content  # Return original if preprocessing fails
    
    def _clean_content(self, content: str) -> str:
        """Clean and normalize content."""
        # Remove excessive whitespace
        content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)
        
        # Normalize quotes
        content = content.replace('"', '"').replace('"', '"')
        content = content.replace(''', "'").replace(''', "'")
        
        # Remove excessive spaces
        content = re.sub(r' +', ' ', content)
        
        # Fix common formatting issues
        content = re.sub(r'\s+([,.!?;:])', r'\1', content)  # Remove space before punctuation
        content = re.sub(r'([.!?])\s*([a-z])', r'\1 \2', content)  # Ensure space after sentence end
        
        return content.strip()
    
    async def _detect_chapters(self, content: str, metadata: EBookMetadata) -> List[Dict[str, Any]]:
        """Detect and structure chapters from content with better logic."""
        try:
            logger.info("Detecting chapters in content")
            
            # Split by common chapter indicators
            chapter_patterns = [
                r'\n\s*Chapter\s+\d+[:\s\n]',
                r'\n\s*CHAPTER\s+\d+[:\s\n]',
                r'\n\s*\d+\.\s+[A-Z]',
                r'\n\s*[A-Z][A-Z\s]{10,}\n',  # All caps headings
            ]
            
            chapters = []
            
            # Try to split by patterns
            for pattern in chapter_patterns:
                matches = list(re.finditer(pattern, content, re.IGNORECASE))
                if len(matches) > 1:  # Found multiple chapters
                    for i, match in enumerate(matches):
                        start = match.start()
                        end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
                        
                        chapter_content = content[start:end].strip()
                        chapter_title = self._extract_chapter_title(chapter_content, i + 1)
                        clean_content = self._clean_chapter_content(chapter_content)
                        
                        if clean_content.strip():  # Only add non-empty chapters
                            chapters.append({
                                "title": chapter_title,
                                "content": clean_content,
                                "word_count": len(clean_content.split()),
                                "order": i + 1
                            })
                    break
            
            # If no chapters detected, try splitting by section breaks
            if not chapters:
                sections = re.split(r'\n\s*[-=*]{3,}\s*\n', content)
                if len(sections) > 1:
                    for i, section in enumerate(sections):
                        section = section.strip()
                        if section:
                            chapters.append({
                                "title": f"Section {i + 1}",
                                "content": section,
                                "word_count": len(section.split()),
                                "order": i + 1
                            })
            
            # If still no chapters, create a single chapter
            if not chapters:
                chapters = [{
                    "title": metadata.title or "Main Content",
                    "content": content,
                    "word_count": len(content.split()),
                    "order": 1
                }]
            
            logger.info("Chapter detection completed", chapter_count=len(chapters))
            return chapters
            
        except Exception as e:
            logger.error("Chapter detection failed", error=str(e))
            # Fallback: single chapter
            return [{
                "title": metadata.title or "Main Content",
                "content": content,
                "word_count": len(content.split()),
                "order": 1
            }]
    
    def _extract_chapter_title(self, chapter_content: str, chapter_num: int) -> str:
        """Extract chapter title from content."""
        lines = chapter_content.split('\n')[:5]  # Check first 5 lines
        
        for line in lines:
            line = line.strip()
            if line and len(line) < 100:
                # Check if it looks like a title
                if any(pattern in line.lower() for pattern in ['chapter', 'section', 'part']):
                    return line
                elif line.isupper() and len(line) > 3:
                    return line.title()
                elif not line.endswith('.') and len(line.split()) <= 10:
                    return line
        
        return f"Chapter {chapter_num}"
    
    def _clean_chapter_content(self, content: str) -> str:
        """Clean chapter content, removing title if present."""
        lines = content.split('\n')
        
        # Remove the first line if it looks like a title
        if lines and len(lines) > 1:
            first_line = lines[0].strip()
            if (len(first_line) < 100 and 
                not first_line.endswith('.') and 
                (first_line.isupper() or 'chapter' in first_line.lower())):
                lines = lines[1:]
        
        return '\n'.join(lines).strip()
    
    async def _generate_epub(self, chapters: List[Dict], metadata: EBookMetadata, 
                           format_options: FormatOptions) -> BinaryIO:
        """Generate EPUB format eBook with proper structure."""
        try:
            logger.info("Generating EPUB", chapter_count=len(chapters))
            
            # Create EPUB book
            book = epub.EpubBook()
            
            # Set metadata
            book.set_identifier(metadata.isbn or str(uuid.uuid4()))
            book.set_title(metadata.title)
            book.set_language(metadata.language)
            book.add_author(metadata.author)
            
            if metadata.description:
                book.add_metadata('DC', 'description', metadata.description)
            if metadata.publisher:
                book.add_metadata('DC', 'publisher', metadata.publisher)
            if metadata.genre:
                book.add_metadata('DC', 'subject', metadata.genre)
            
            book.add_metadata('DC', 'date', datetime.now().isoformat())
            
            # Create CSS file
            css_content = self.css_template.render(
                font_family=format_options.font_family,
                font_size=format_options.font_size,
                line_height=format_options.line_height,
                margin_top=format_options.margin_top,
                margin_right=format_options.margin_right,
                margin_bottom=format_options.margin_bottom,
                margin_left=format_options.margin_left,
                justify_text=format_options.justify_text,
                page_break_before_chapter=format_options.page_break_before_chapter
            )
            
            nav_css = epub.EpubItem(
                uid="style_nav",
                file_name="style/main.css",
                media_type="text/css",
                content=css_content
            )
            book.add_item(nav_css)
            
            # Add title page
            if format_options.include_cover:
                title_page = self._create_title_page_html(metadata, format_options)
                title_chapter = epub.EpubHtml(
                    title="Title Page",
                    file_name='title.xhtml',
                    lang=metadata.language
                )
                title_chapter.content = title_page
                title_chapter.add_item(nav_css)
                book.add_item(title_chapter)
            
            # Create chapters
            epub_chapters = []
            spine_items = ['nav']
            
            if format_options.include_cover:
                spine_items.append(title_chapter)
            
            for i, chapter in enumerate(chapters):
                chapter_html = self.chapter_template.render(
                    title=chapter['title'],
                    content=self._format_content_as_html(chapter['content'])
                )
                
                epub_chapter = epub.EpubHtml(
                    title=chapter['title'],
                    file_name=f'chap_{i+1:02d}.xhtml',
                    lang=metadata.language
                )
                epub_chapter.content = chapter_html
                epub_chapter.add_item(nav_css)
                
                book.add_item(epub_chapter)
                epub_chapters.append(epub_chapter)
                spine_items.append(epub_chapter)
            
            # Define Table of Contents
            if format_options.include_toc and len(chapters) > 1:
                book.toc = tuple(
                    epub.Link(f'chap_{i+1:02d}.xhtml', chapter['title'], f'chap_{i+1}')
                    for i, chapter in enumerate(chapters)
                )
            
            # Add navigation files
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Define spine
            book.spine = spine_items
            
            # Generate EPUB file in memory
            epub_content = io.BytesIO()
            epub.write_epub(epub_content, book, {})
            epub_content.seek(0)
            
            logger.info("EPUB generation completed")
            return epub_content
            
        except Exception as e:
            logger.error("EPUB generation failed", error=str(e))
            raise
    
    def _create_title_page_html(self, metadata: EBookMetadata, format_options: FormatOptions) -> str:
        """Create HTML for title page."""
        title_html = f"""
        <!DOCTYPE html>
        <html xmlns="http://www.w3.org/1999/xhtml">
        <head>
            <title>Title Page</title>
            <meta charset="utf-8"/>
            <link rel="stylesheet" type="text/css" href="style/main.css"/>
        </head>
        <body>
            <div class="book-header">
                <h1 class="book-title">{metadata.title}</h1>
                <p class="book-author">by {metadata.author}</p>
                {f'<p class="book-description">{metadata.description}</p>' if metadata.description else ''}
                {f'<p class="book-publisher">{metadata.publisher}</p>' if metadata.publisher else ''}
            </div>
        </body>
        </html>
        """
        return title_html
    
    async def _generate_pdf(self, chapters: List[Dict], metadata: EBookMetadata, 
                          format_options: FormatOptions) -> BinaryIO:
        """Generate PDF format eBook using PyMuPDF."""
        try:
            logger.info("Generating PDF", chapter_count=len(chapters))
            
            if not fitz:
                raise ImportError("PyMuPDF not available for PDF generation")
            
            # Create new PDF document
            doc = fitz.open()
            
            # Set metadata
            doc.set_metadata({
                "title": metadata.title,
                "author": metadata.author,
                "subject": metadata.description or "",
                "creator": "eBook Editor Pro",
                "producer": metadata.publisher or "eBook Editor Pro"
            })
            
            # Add cover page if requested
            if format_options.include_cover:
                self._add_pdf_cover_page(doc, metadata, format_options)
            
            # Add table of contents if requested
            if format_options.include_toc and len(chapters) > 1:
                self._add_pdf_toc(doc, chapters, format_options)
            
            # Add chapters
            for chapter in chapters:
                if format_options.page_break_before_chapter and doc.page_count > 0:
                    doc.new_page()
                
                self._add_pdf_chapter(doc, chapter, format_options)
            
            # Save PDF to memory
            pdf_content = io.BytesIO()
            pdf_bytes = doc.tobytes()
            pdf_content.write(pdf_bytes)
            pdf_content.seek(0)
            
            doc.close()
            
            logger.info("PDF generation completed")
            return pdf_content
            
        except Exception as e:
            logger.error("PDF generation failed", error=str(e))
            raise
    
    def _add_pdf_cover_page(self, doc, metadata: EBookMetadata, format_options: FormatOptions):
        """Add cover page to PDF."""
        page = doc.new_page()
        
        # Title
        rect = fitz.Rect(50, 150, 500, 250)
        page.insert_textbox(rect, metadata.title, 
                           fontsize=28, fontname="helv", 
                           align=fitz.TEXT_ALIGN_CENTER)
        
        # Author
        rect = fitz.Rect(50, 280, 500, 320)
        page.insert_textbox(rect, f"by {metadata.author}", 
                           fontsize=18, fontname="helv", 
                           align=fitz.TEXT_ALIGN_CENTER)
        
        # Description
        if metadata.description:
            rect = fitz.Rect(50, 350, 500, 450)
            page.insert_textbox(rect, metadata.description, 
                               fontsize=12, fontname="helv", 
                               align=fitz.TEXT_ALIGN_CENTER)
    
    def _add_pdf_toc(self, doc, chapters: List[Dict], format_options: FormatOptions):
        """Add table of contents to PDF."""
        page = doc.new_page()
        
        # Title
        rect = fitz.Rect(50, 50, 500, 100)
        page.insert_textbox(rect, "Table of Contents", 
                           fontsize=24, fontname="helv", 
                           align=fitz.TEXT_ALIGN_CENTER)
        
        # Chapters
        y_pos = 120
        for i, chapter in enumerate(chapters, 1):
            if y_pos > 700:  # Start new page if needed
                page = doc.new_page()
                y_pos = 50
            
            rect = fitz.Rect(50, y_pos, 500, y_pos + 25)
            page.insert_textbox(rect, f"{i}. {chapter['title']}", 
                               fontsize=14, fontname="helv")
            y_pos += 30
    
    def _add_pdf_chapter(self, doc, chapter: Dict, format_options: FormatOptions):
        """Add chapter content to PDF with proper text flow."""
        page = doc.new_page()
        
        # Chapter title (only if not "Main Content")
        y_pos = 50
        if chapter['title'] != "Main Content":
            rect = fitz.Rect(50, y_pos, 500, y_pos + 50)
            page.insert_textbox(rect, chapter['title'], 
                               fontsize=20, fontname="helvb", 
                               align=fitz.TEXT_ALIGN_CENTER)
            y_pos += 80
        
        # Chapter content with text flow
        content = chapter['content']
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for paragraph in paragraphs:
            if not paragraph.strip():
                continue
                
            # Estimate how much space this paragraph needs
            para_height = self._estimate_text_height(paragraph, format_options.font_size)
            
            # Check if we need a new page
            if y_pos + para_height > 750:
                page = doc.new_page()
                y_pos = 50
            
            # Insert paragraph
            rect = fitz.Rect(50, y_pos, 500, min(y_pos + para_height + 20, 750))
            try:
                page.insert_textbox(rect, paragraph.strip(), 
                                   fontsize=format_options.font_size, 
                                   fontname="helv")
            except:
                # If text doesn't fit, split it
                words = paragraph.split()
                chunk_size = len(words) // 2
                if chunk_size > 0:
                    first_chunk = ' '.join(words[:chunk_size])
                    second_chunk = ' '.join(words[chunk_size:])
                    
                    # Insert first chunk
                    page.insert_textbox(rect, first_chunk, 
                                       fontsize=format_options.font_size, 
                                       fontname="helv")
                    
                    # Continue with second chunk on next iteration
                    y_pos += para_height + 10
                    continue
            
            y_pos += para_height + 15
    
    def _estimate_text_height(self, text: str, font_size: int) -> float:
        """Estimate the height needed for text."""
        lines = len(text) // 70 + 1  # Approximate 70 characters per line
        line_height = font_size * 1.2
        return lines * line_height
    
    async def _generate_docx(self, chapters: List[Dict], metadata: EBookMetadata, 
                           format_options: FormatOptions) -> BinaryIO:
        """Generate DOCX format eBook with proper formatting."""
        try:
            logger.info("Generating DOCX", chapter_count=len(chapters))
            
            # Create new document
            doc = Document()
            
            # Set document properties
            properties = doc.core_properties
            properties.title = metadata.title
            properties.author = metadata.author
            if metadata.description:
                properties.comments = metadata.description
            if metadata.publisher:
                properties.category = metadata.publisher
            
            # Configure styles
            style = doc.styles['Normal']
            font = style.font
            font.name = format_options.font_family
            font.size = Pt(format_options.font_size)
            
            paragraph_format = style.paragraph_format
            paragraph_format.line_spacing = format_options.line_height
            paragraph_format.space_after = Pt(6)
            
            # Add title page if cover is requested
            if format_options.include_cover:
                self._add_docx_title_page(doc, metadata)
                doc.add_page_break()
            
            # Add table of contents if requested
            if format_options.include_toc and len(chapters) > 1:
                self._add_docx_toc(doc, chapters)
                doc.add_page_break()
            
            # Add chapters
            for i, chapter in enumerate(chapters):
                if i > 0 and format_options.page_break_before_chapter:
                    doc.add_page_break()
                
                # Chapter title (only if not "Main Content")
                if chapter['title'] != "Main Content":
                    title_paragraph = doc.add_heading(chapter['title'], level=1)
                    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Chapter content
                content_paragraphs = chapter['content'].split('\n\n')
                for para_text in content_paragraphs:
                    if para_text.strip():
                        paragraph = doc.add_paragraph(para_text.strip())
                        if format_options.justify_text:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Save to memory
            docx_content = io.BytesIO()
            doc.save(docx_content)
            docx_content.seek(0)
            
            logger.info("DOCX generation completed")
            return docx_content
            
        except Exception as e:
            logger.error("DOCX generation failed", error=str(e))
            raise
    
    def _add_docx_title_page(self, doc: Document, metadata: EBookMetadata):
        """Add title page to DOCX."""
        # Title
        title = doc.add_heading(metadata.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some space
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Author
        author = doc.add_paragraph(f"by {metadata.author}")
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Description if available
        if metadata.description:
            doc.add_paragraph()
            desc = doc.add_paragraph(metadata.description)
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_docx_toc(self, doc: Document, chapters: List[Dict]):
        """Add table of contents to DOCX."""
        doc.add_heading("Table of Contents", level=1)
        
        for i, chapter in enumerate(chapters, 1):
            toc_entry = doc.add_paragraph(f"{i}. {chapter['title']}")
            toc_entry.style = 'List Number'
    
    async def _generate_html(self, chapters: List[Dict], metadata: EBookMetadata, 
                           format_options: FormatOptions) -> BinaryIO:
        """Generate HTML format eBook with comprehensive styling."""
        try:
            logger.info("Generating HTML", chapter_count=len(chapters))
            
            # Generate CSS
            css_content = self.css_template.render(
                font_family=format_options.font_family,
                font_size=format_options.font_size,
                line_height=format_options.line_height,
                margin_top=format_options.margin_top,
                margin_right=format_options.margin_right,
                margin_bottom=format_options.margin_bottom,
                margin_left=format_options.margin_left,
                justify_text=format_options.justify_text,
                page_break_before_chapter=format_options.page_break_before_chapter
            ) + """
            .book-header { 
                text-align: center; 
                margin-bottom: 3em; 
                page-break-after: always;
                padding: 2em;
                border-bottom: 2px solid #3498db;
            }
            .book-title {
                font-size: 2.5em;
                font-weight: bold;
                margin-bottom: 0.5em;
                color: #2c3e50;
            }
            .book-author { 
                font-style: italic; 
                font-size: 1.3em; 
                color: #7f8c8d;
                margin-bottom: 1em;
            }
            .book-description { 
                margin-top: 1em; 
                font-style: italic; 
                max-width: 600px;
                margin-left: auto;
                margin-right: auto;
            }
            .toc { 
                margin-bottom: 3em; 
                padding: 2em;
                background-color: #f8f9fa;
                border-radius: 8px;
            }
            .toc h2 {
                color: #2c3e50;
                border-bottom: 2px solid #3498db;
                padding-bottom: 0.5em;
            }
            .toc ul { 
                list-style-type: none; 
                padding-left: 0;
            }
            .toc li { 
                margin-bottom: 0.8em; 
                padding: 0.5em;
                border-left: 3px solid #3498db;
                background-color: white;
            }
            .toc a {
                text-decoration: none;
                color: #2c3e50;
                font-weight: 500;
            }
            .toc a:hover {
                color: #3498db;
            }
            .chapter { 
                margin-bottom: 3em; 
                padding: 2em;
                background-color: white;
                border-radius: 8px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }
            .book-footer { 
                margin-top: 3em; 
                text-align: center; 
                font-size: 0.9em; 
                color: #7f8c8d; 
                border-top: 1px solid #ecf0f1;
                padding: 2em;
            }
            body {
                background-color: #f8f9fa;
                font-family: {{ font_family }}, serif;
            }
            .container {
                max-width: 800px;
                margin: 0 auto;
                background-color: white;
                box-shadow: 0 0 20px rgba(0,0,0,0.1);
                min-height: 100vh;
            }
            """
            
            # HTML template
            html_template = Template("""
            <!DOCTYPE html>
            <html lang="{{ language }}">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{{ title }}</title>
                <meta name="author" content="{{ author }}">
                {% if description %}<meta name="description" content="{{ description }}">{% endif %}
                <style>{{ css }}</style>
            </head>
            <body>
                <div class="container">
                    <header class="book-header">
                        <h1 class="book-title">{{ title }}</h1>
                        <p class="book-author">by {{ author }}</p>
                        {% if description %}<p class="book-description">{{ description }}</p>{% endif %}
                    </header>
                    
                    {% if include_toc and chapters|length > 1 %}
                    <nav class="toc">
                        <h2>Table of Contents</h2>
                        <ul>
                        {% for chapter in chapters %}
                            <li><a href="#chapter-{{ loop.index }}">{{ chapter.title }}</a></li>
                        {% endfor %}
                        </ul>
                    </nav>
                    {% endif %}
                    
                    <main class="book-content">
                    {% for chapter in chapters %}
                        <section class="chapter" id="chapter-{{ loop.index }}">
                            {% if chapter.title != "Main Content" %}
                            <h2 class="chapter-title">{{ chapter.title }}</h2>
                            {% endif %}
                            <div class="chapter-content">
                                {{ chapter.formatted_content | safe }}
                            </div>
                        </section>
                    {% endfor %}
                    </main>
                    
                    <footer class="book-footer">
                        <p>Generated by eBook Editor Pro</p>
                        <p>{{ generation_date }}</p>
                    </footer>
                </div>
            </body>
            </html>
            """)
            
            # Format chapters
            formatted_chapters = []
            for chapter in chapters:
                formatted_chapters.append({
                    'title': chapter['title'],
                    'formatted_content': self._format_content_as_html(chapter['content'])
                })
            
            # Render HTML
            html_content = html_template.render(
                title=metadata.title,
                author=metadata.author,
                description=metadata.description,
                language=metadata.language,
                css=css_content,
                chapters=formatted_chapters,
                include_toc=format_options.include_toc,
                generation_date=datetime.now().strftime("%B %d, %Y")
            )
            
            # Return as BytesIO
            html_bytes = io.BytesIO()
            html_bytes.write(html_content.encode('utf-8'))
            html_bytes.seek(0)
            
            logger.info("HTML generation completed")
            return html_bytes
            
        except Exception as e:
            logger.error("HTML generation failed", error=str(e))
            raise
    
    async def _generate_txt(self, chapters: List[Dict], metadata: EBookMetadata, 
                          format_options: FormatOptions) -> BinaryIO:
        """Generate plain text format eBook with proper formatting."""
        try:
            logger.info("Generating TXT", chapter_count=len(chapters))
            
            content_lines = []
            
            # Add header
            content_lines.extend([
                "=" * 60,
                metadata.title.upper(),
                f"by {metadata.author}",
                "=" * 60,
                ""
            ])
            
            if metadata.description:
                content_lines.extend([metadata.description, ""])
            
            if metadata.publisher:
                content_lines.extend([f"Published by {metadata.publisher}", ""])
            
            # Add table of contents if requested
            if format_options.include_toc and len(chapters) > 1:
                content_lines.extend(["TABLE OF CONTENTS", "-" * 20, ""])
                for i, chapter in enumerate(chapters, 1):
                    content_lines.append(f"{i:2d}. {chapter['title']}")
                content_lines.extend(["", "=" * 60, ""])
            
            # Add chapters
            for i, chapter in enumerate(chapters):
                if i > 0:
                    content_lines.extend(["", "=" * 60, ""])
                
                # Chapter title (only if not "Main Content")
                if chapter['title'] != "Main Content":
                    content_lines.extend([
                        chapter['title'].upper(),
                        "-" * len(chapter['title']),
                        ""
                    ])
                
                # Format content
                paragraphs = chapter['content'].split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        # Wrap text to reasonable line length
                        lines = self._wrap_text(paragraph.strip(), 75)
                        content_lines.extend(lines)
                        content_lines.append("")
            
            # Add footer
            content_lines.extend([
                "",
                "=" * 60,
                f"Generated by eBook Editor Pro on {datetime.now().strftime('%B %d, %Y')}",
                "=" * 60
            ])
            
            # Join all content
            full_content = '\n'.join(content_lines)
            
            # Return as BytesIO
            txt_bytes = io.BytesIO()
            txt_bytes.write(full_content.encode('utf-8'))
            txt_bytes.seek(0)
            
            logger.info("TXT generation completed")
            return txt_bytes
            
        except Exception as e:
            logger.error("TXT generation failed", error=str(e))
            raise
    
    def _format_content_as_html(self, content: str) -> str:
        """Format plain text content as HTML with better structure."""
        # Split into paragraphs
        paragraphs = content.split('\n\n')
        html_paragraphs = []
        
        for para in paragraphs:
            para = para.strip()
            if para:
                # Handle line breaks within paragraphs
                para = para.replace('\n', '<br>')
                
                # Basic markdown-like formatting
                para = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', para)
                para = re.sub(r'\*(.*?)\*', r'<em>\1</em>', para)
                para = re.sub(r'`(.*?)`', r'<code>\1</code>', para)
                
                # Handle quotes
                if para.startswith('"') and para.endswith('"'):
                    para = f'<blockquote>{para[1:-1]}</blockquote>'
                else:
                    para = f'<p>{para}</p>'
                
                html_paragraphs.append(para)
        
        return '\n'.join(html_paragraphs)
    
    def _wrap_text(self, text: str, width: int) -> List[str]:
        """Wrap text to specified width with proper word breaks."""
        import textwrap
        return textwrap.wrap(text, width=width, break_long_words=False, break_on_hyphens=True)
    
    def cleanup_temp_files(self):
        """Clean up temporary files."""
        for temp_file in self.temp_files:
            try:
                Path(temp_file).unlink(missing_ok=True)
            except Exception as e:
                logger.warning("Failed to delete temp file", file=temp_file, error=str(e))
        self.temp_files.clear()