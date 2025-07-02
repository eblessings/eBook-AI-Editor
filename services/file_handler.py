"""
Fixed File Handler Service for the eBook Editor.
Handles file uploads, text extraction from various formats, and file processing.
Supports DOCX, PDF, EPUB, TXT, HTML, and Markdown files with improved error handling.
"""

import asyncio
import hashlib
import io
import mimetypes
import os
import tempfile
import uuid
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Optional, Any, BinaryIO, Union
import aiofiles
import chardet
from urllib.parse import urlparse

# File processing libraries
try:
    import mammoth
except ImportError:
    mammoth = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from ebooklib import epub
    import ebooklib
except ImportError:
    epub = None
    ebooklib = None

from bs4 import BeautifulSoup

try:
    import docx2txt
    from docx import Document
except ImportError:
    docx2txt = None
    Document = None

try:
    import magic
except ImportError:
    magic = None

import structlog
import re

from config import Settings

logger = structlog.get_logger()


class FileHandler:
    """Enhanced file handler with multi-format support and better error handling."""
    
    def __init__(self, settings: Settings):
        self.settings = settings
        self.temp_files = []
        self.processing_cache = {}
        
        # Ensure directories exist
        Path(self.settings.TEMP_DIR).mkdir(parents=True, exist_ok=True)
        Path(self.settings.UPLOAD_DIR).mkdir(parents=True, exist_ok=True)
        Path(self.settings.EXPORT_DIR).mkdir(parents=True, exist_ok=True)
    
    async def extract_text_from_file(self, file) -> Dict[str, Any]:
        """Extract text content from uploaded file with enhanced error handling."""
        try:
            logger.info("Extracting text from file", 
                       filename=file.filename, 
                       content_type=file.content_type)
            
            # Read file content
            content = await file.read()
            file_size = len(content)
            
            # Validate file size
            if file_size > self.settings.MAX_FILE_SIZE:
                raise ValueError(f"File too large: {file_size} bytes")
            
            # Detect actual content type
            detected_type = self._detect_content_type(content, file.filename)
            logger.debug("Content type detected", 
                        original=file.content_type, 
                        detected=detected_type)
            
            # Generate file ID and save temporarily
            file_id = str(uuid.uuid4())
            temp_path = await self._save_temp_file(content, file_id, file.filename)
            
            # Extract text based on file type
            extracted_data = await self._extract_by_type(
                content, temp_path, detected_type, file.filename
            )
            
            # Calculate additional metadata
            word_count = len(extracted_data['text'].split())
            char_count = len(extracted_data['text'])
            
            # Clean up temp file
            await self._cleanup_temp_file(temp_path)
            
            result = {
                'file_id': file_id,
                'filename': file.filename,
                'original_type': file.content_type,
                'detected_type': detected_type,
                'file_size': file_size,
                'text': extracted_data['text'],
                'word_count': word_count,
                'character_count': char_count,
                'metadata': extracted_data.get('metadata', {}),
                'structure': extracted_data.get('structure', {}),
                'images': extracted_data.get('images', []),
                'extraction_method': extracted_data.get('method', 'unknown')
            }
            
            logger.info("Text extraction completed", 
                       file_id=file_id, 
                       word_count=word_count)
            
            return result
            
        except Exception as e:
            logger.error("Text extraction failed", 
                        filename=getattr(file, 'filename', 'unknown'), 
                        error=str(e))
            raise
    
    def _detect_content_type(self, content: bytes, filename: str) -> str:
        """Detect actual content type using multiple methods."""
        try:
            # Check file signatures first
            if content.startswith(b'PK\x03\x04'):
                # ZIP-based format
                if self._is_docx(content):
                    return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                elif b'epub' in content[:1024].lower():
                    return 'application/epub+zip'
                else:
                    return 'application/zip'
            elif content.startswith(b'%PDF'):
                return 'application/pdf'
            elif content.startswith((b'\xff\xfe', b'\xfe\xff', b'\xef\xbb\xbf')):
                return 'text/plain'  # UTF encoded text
            elif content.startswith(b'<!DOCTYPE html') or content.startswith(b'<html'):
                return 'text/html'
        except:
            pass
        
        # Try python-magic if available
        try:
            if magic:
                detected_mime = magic.from_buffer(content, mime=True)
                if detected_mime and detected_mime != 'application/octet-stream':
                    return detected_mime
        except:
            pass
        
        # Fallback to filename extension
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type:
            return mime_type
        
        # Default fallback
        return 'application/octet-stream'
    
    def _is_docx(self, content: bytes) -> bool:
        """Check if ZIP content is a DOCX file."""
        try:
            with zipfile.ZipFile(io.BytesIO(content), 'r') as zip_file:
                return 'word/document.xml' in zip_file.namelist()
        except:
            return False
    
    async def _save_temp_file(self, content: bytes, file_id: str, filename: str) -> str:
        """Save content to temporary file."""
        temp_path = os.path.join(
            self.settings.TEMP_DIR, 
            f"{file_id}_{filename}"
        )
        
        async with aiofiles.open(temp_path, 'wb') as f:
            await f.write(content)
        
        self.temp_files.append(temp_path)
        return temp_path
    
    async def _cleanup_temp_file(self, temp_path: str):
        """Clean up temporary file."""
        try:
            if temp_path in self.temp_files:
                self.temp_files.remove(temp_path)
            await aiofiles.os.remove(temp_path)
        except Exception as e:
            logger.warning("Failed to cleanup temp file", path=temp_path, error=str(e))
    
    async def _extract_by_type(self, content: bytes, temp_path: str, 
                              content_type: str, filename: str) -> Dict[str, Any]:
        """Extract text based on content type with enhanced error handling."""
        try:
            # Route to appropriate extraction method
            if content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return await self._extract_from_docx(temp_path, content)
            elif content_type == 'application/pdf':
                return await self._extract_from_pdf(temp_path)
            elif content_type == 'application/epub+zip':
                return await self._extract_from_epub(temp_path)
            elif content_type in ['text/plain', 'text/markdown']:
                return await self._extract_from_text(content)
            elif content_type == 'text/html':
                return await self._extract_from_html(content)
            else:
                # Try as text
                logger.warning("Unknown content type, trying as text", 
                              content_type=content_type)
                return await self._extract_from_text(content)
                
        except Exception as e:
            logger.error("Type-specific extraction failed", 
                        content_type=content_type, 
                        error=str(e))
            # Fallback to basic text extraction
            return await self._extract_from_text(content)
    
    async def _extract_from_docx(self, file_path: str, content: bytes = None) -> Dict[str, Any]:
        """Extract text from DOCX file with multiple fallback methods."""
        try:
            logger.debug("Extracting from DOCX")
            
            # Method 1: Try mammoth for best formatting
            if mammoth:
                try:
                    with open(file_path, 'rb') as docx_file:
                        result = mammoth.convert_to_html(docx_file)
                        html_content = result.value
                        
                        # Convert HTML to plain text while preserving structure
                        soup = BeautifulSoup(html_content, 'html.parser')
                        text_content = soup.get_text('\n\n', strip=True)
                        
                        if text_content and len(text_content.strip()) > 0:
                            logger.debug("Successfully extracted using mammoth")
                            return {
                                'text': text_content,
                                'metadata': {},
                                'structure': {'extraction_method': 'mammoth'},
                                'method': 'mammoth',
                                'html_content': html_content
                            }
                except Exception as e:
                    logger.warning("Mammoth extraction failed", error=str(e))
            
            # Method 2: Try python-docx for metadata
            if Document:
                try:
                    doc = Document(file_path)
                    
                    # Extract text from paragraphs
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text.strip())
                    
                    text_content = '\n\n'.join(paragraphs)
                    
                    if text_content and len(text_content.strip()) > 0:
                        # Extract metadata
                        metadata = {
                            'title': doc.core_properties.title or '',
                            'author': doc.core_properties.author or '',
                            'subject': doc.core_properties.subject or '',
                            'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                            'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
                        }
                        
                        structure = {
                            'paragraphs': len(paragraphs),
                            'tables': len(doc.tables),
                            'sections': len(doc.sections),
                            'extraction_method': 'python-docx'
                        }
                        
                        logger.debug("Successfully extracted using python-docx")
                        return {
                            'text': text_content,
                            'metadata': metadata,
                            'structure': structure,
                            'method': 'python-docx'
                        }
                except Exception as e:
                    logger.warning("python-docx extraction failed", error=str(e))
            
            # Method 3: Try docx2txt as fallback
            if docx2txt:
                try:
                    text_content = docx2txt.process(file_path)
                    if text_content and len(text_content.strip()) > 0:
                        logger.debug("Successfully extracted using docx2txt")
                        return {
                            'text': text_content,
                            'metadata': {},
                            'structure': {'extraction_method': 'docx2txt'},
                            'method': 'docx2txt'
                        }
                except Exception as e:
                    logger.warning("docx2txt extraction failed", error=str(e))
            
            # Method 4: Manual XML parsing as last resort
            try:
                return await self._extract_docx_manually(file_path, content)
            except Exception as e:
                logger.error("Manual DOCX extraction failed", error=str(e))
                raise ValueError(f"Failed to extract text from DOCX: All methods failed")
                
        except Exception as e:
            logger.error("DOCX extraction completely failed", error=str(e))
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_docx_manually(self, file_path: str, content: bytes = None) -> Dict[str, Any]:
        """Manually extract text from DOCX by parsing XML."""
        try:
            logger.debug("Attempting manual DOCX extraction")
            
            if content is None:
                with open(file_path, 'rb') as f:
                    content = f.read()
            
            with zipfile.ZipFile(io.BytesIO(content), 'r') as zip_file:
                # Read the main document
                document_xml = zip_file.read('word/document.xml')
                
                # Parse XML
                root = ET.fromstring(document_xml)
                
                # Define namespaces
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                # Extract text from paragraphs
                paragraphs = []
                for para in root.findall('.//w:p', namespaces):
                    para_text = ''
                    for text_elem in para.findall('.//w:t', namespaces):
                        if text_elem.text:
                            para_text += text_elem.text
                    
                    if para_text.strip():
                        paragraphs.append(para_text.strip())
                
                text_content = '\n\n'.join(paragraphs)
                
                if text_content and len(text_content.strip()) > 0:
                    logger.debug("Successfully extracted using manual XML parsing")
                    return {
                        'text': text_content,
                        'metadata': {},
                        'structure': {
                            'paragraphs': len(paragraphs),
                            'extraction_method': 'manual_xml'
                        },
                        'method': 'manual_xml'
                    }
                else:
                    raise ValueError("No text content found in DOCX")
                    
        except Exception as e:
            logger.error("Manual DOCX extraction failed", error=str(e))
            raise
    
    async def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file with fallback."""
        try:
            logger.debug("Extracting from PDF")
            
            if not fitz:
                raise ImportError("PyMuPDF not available for PDF processing")
            
            # Open PDF with PyMuPDF
            doc = fitz.open(file_path)
            
            text_content = []
            metadata = doc.metadata
            images = []
            structure = {
                'pages': doc.page_count,
                'has_toc': len(doc.get_toc()) > 0,
                'toc_entries': len(doc.get_toc())
            }
            
            # Extract text from each page
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                if page_text.strip():
                    text_content.append(page_text)
                
                # Extract images (metadata only)
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    images.append({
                        'page': page_num + 1,
                        'index': img_index,
                        'xref': img[0]
                    })
            
            doc.close()
            
            # Join all text with page breaks
            full_text = '\n\n--- Page Break ---\n\n'.join(text_content)
            
            return {
                'text': full_text,
                'metadata': {
                    'title': metadata.get('title', ''),
                    'author': metadata.get('author', ''),
                    'subject': metadata.get('subject', ''),
                    'creator': metadata.get('creator', ''),
                    'producer': metadata.get('producer', ''),
                    'creation_date': metadata.get('creationDate', ''),
                    'modification_date': metadata.get('modDate', '')
                },
                'structure': structure,
                'images': images,
                'method': 'pymupdf'
            }
            
        except Exception as e:
            logger.error("PDF extraction failed", error=str(e))
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_from_epub(self, file_path: str) -> Dict[str, Any]:
        """Extract text from EPUB file."""
        try:
            logger.debug("Extracting from EPUB")
            
            if not epub or not ebooklib:
                raise ImportError("ebooklib not available for EPUB processing")
            
            # Open EPUB with ebooklib
            book = epub.read_epub(file_path)
            
            # Extract metadata
            metadata = {
                'title': book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else '',
                'author': book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else '',
                'language': book.get_metadata('DC', 'language')[0][0] if book.get_metadata('DC', 'language') else '',
                'publisher': book.get_metadata('DC', 'publisher')[0][0] if book.get_metadata('DC', 'publisher') else '',
                'description': book.get_metadata('DC', 'description')[0][0] if book.get_metadata('DC', 'description') else '',
                'identifier': book.get_metadata('DC', 'identifier')[0][0] if book.get_metadata('DC', 'identifier') else ''
            }
            
            # Extract text from all document items
            text_content = []
            images = []
            chapter_count = 0
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    # Parse HTML content
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    chapter_text = soup.get_text('\n', strip=True)
                    
                    if chapter_text.strip():
                        text_content.append(chapter_text)
                        chapter_count += 1
                        
                elif item.get_type() == ebooklib.ITEM_IMAGE:
                    images.append({
                        'filename': item.get_name(),
                        'media_type': item.get_type()
                    })
            
            # Join chapters
            full_text = '\n\n--- Chapter Break ---\n\n'.join(text_content)
            
            structure = {
                'chapters': chapter_count,
                'images': len(images),
                'has_toc': len(book.toc) > 0
            }
            
            return {
                'text': full_text,
                'metadata': metadata,
                'structure': structure,
                'images': images,
                'method': 'ebooklib'
            }
            
        except Exception as e:
            logger.error("EPUB extraction failed", error=str(e))
            raise ValueError(f"Failed to extract text from EPUB: {str(e)}")
    
    async def _extract_from_text(self, content: bytes) -> Dict[str, Any]:
        """Extract text from plain text file with robust encoding detection."""
        try:
            logger.debug("Extracting from text file")
            
            # Detect encoding with multiple methods
            encoding = self._detect_encoding(content)
            
            logger.debug("Encoding detected", encoding=encoding)
            
            # Decode text
            text = self._decode_with_fallback(content, encoding)
            
            # Clean and normalize text
            text = self._normalize_text(text)
            
            # Basic structure analysis
            lines = text.split('\n')
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            
            structure = {
                'lines': len(lines),
                'paragraphs': len(paragraphs),
                'encoding': encoding,
                'encoding_confidence': self._get_encoding_confidence(content, encoding)
            }
            
            return {
                'text': text,
                'metadata': {'encoding': encoding},
                'structure': structure,
                'method': 'chardet+decode'
            }
            
        except Exception as e:
            logger.error("Text extraction failed", error=str(e))
            raise ValueError(f"Failed to extract text: {str(e)}")
    
    def _detect_encoding(self, content: bytes) -> str:
        """Detect encoding using multiple methods."""
        # Method 1: Check for BOM
        if content.startswith(b'\xef\xbb\xbf'):
            return 'utf-8-sig'
        elif content.startswith(b'\xff\xfe'):
            return 'utf-16-le'
        elif content.startswith(b'\xfe\xff'):
            return 'utf-16-be'
        
        # Method 2: Use chardet
        try:
            detected = chardet.detect(content)
            if detected and detected.get('confidence', 0) > 0.7:
                return detected['encoding']
        except:
            pass
        
        # Method 3: Try common encodings
        common_encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252', 'iso-8859-1']
        for encoding in common_encodings:
            try:
                content.decode(encoding)
                return encoding
            except UnicodeDecodeError:
                continue
        
        # Fallback
        return 'utf-8'
    
    def _decode_with_fallback(self, content: bytes, primary_encoding: str) -> str:
        """Decode content with fallback encodings."""
        encodings_to_try = [primary_encoding, 'utf-8', 'latin1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings_to_try:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Last resort: decode with errors ignored
        return content.decode('utf-8', errors='ignore')
    
    def _get_encoding_confidence(self, content: bytes, encoding: str) -> float:
        """Get confidence score for encoding detection."""
        try:
            detected = chardet.detect(content)
            if detected and detected.get('encoding', '').lower() == encoding.lower():
                return detected.get('confidence', 0.5)
        except:
            pass
        return 0.5
    
    def _normalize_text(self, text: str) -> str:
        """Normalize text content."""
        # Replace various whitespace characters
        text = re.sub(r'\r\n', '\n', text)  # Windows line endings
        text = re.sub(r'\r', '\n', text)    # Old Mac line endings
        
        # Normalize unicode characters
        import unicodedata
        text = unicodedata.normalize('NFKC', text)
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Multiple blank lines
        text = re.sub(r' +', ' ', text)  # Multiple spaces
        
        # Remove control characters except newlines and tabs
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        
        return text.strip()
    
    async def _extract_from_html(self, content: bytes) -> Dict[str, Any]:
        """Extract text from HTML file with improved parsing."""
        try:
            logger.debug("Extracting from HTML")
            
            # Detect encoding
            encoding = self._detect_encoding(content)
            
            # Decode HTML
            html_text = self._decode_with_fallback(content, encoding)
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(html_text, 'html.parser')
            
            # Extract metadata from meta tags
            metadata = {}
            title_tag = soup.find('title')
            if title_tag:
                metadata['title'] = title_tag.get_text().strip()
            
            # Look for common meta tags
            for meta in soup.find_all('meta'):
                name = meta.get('name', '').lower()
                content_attr = meta.get('content', '')
                
                if name in ['author', 'description', 'keywords']:
                    metadata[name] = content_attr
            
            # Extract text content
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Get text with better formatting
            text = soup.get_text('\n', strip=True)
            text = self._normalize_text(text)
            
            # Structure analysis
            structure = {
                'headings': len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])),
                'paragraphs': len(soup.find_all('p')),
                'links': len(soup.find_all('a')),
                'images': len(soup.find_all('img')),
                'encoding': encoding
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'structure': structure,
                'method': 'beautifulsoup',
                'html_content': html_text
            }
            
        except Exception as e:
            logger.error("HTML extraction failed", error=str(e))
            raise ValueError(f"Failed to extract text from HTML: {str(e)}")
    
    async def start_batch_processing(self, file) -> str:
        """Start batch processing for a file."""
        try:
            task_id = str(uuid.uuid4())
            
            # In a real implementation, this would queue the file for background processing
            logger.info("Starting batch processing", task_id=task_id, filename=file.filename)
            
            # For now, just return the task ID
            return task_id
            
        except Exception as e:
            logger.error("Failed to start batch processing", error=str(e))
            raise
    
    async def validate_file(self, file, max_size: Optional[int] = None) -> Dict[str, Any]:
        """Validate uploaded file with comprehensive checks."""
        try:
            # Check file size
            content = await file.read()
            file_size = len(content)
            
            max_allowed = max_size or self.settings.MAX_FILE_SIZE
            if file_size > max_allowed:
                return {
                    'valid': False,
                    'error': f'File too large: {file_size} bytes (max: {max_allowed})'
                }
            
            # Check for empty file
            if file_size == 0:
                return {
                    'valid': False,
                    'error': 'File is empty'
                }
            
            # Check content type
            detected_type = self._detect_content_type(content, file.filename)
            if detected_type not in self.settings.SUPPORTED_UPLOAD_TYPES:
                return {
                    'valid': False,
                    'error': f'Unsupported file type: {detected_type}'
                }
            
            # Check for malicious content (basic checks)
            if self._is_potentially_malicious(content, file.filename):
                return {
                    'valid': False,
                    'error': 'File failed security checks'
                }
            
            # Try to extract a small sample to verify file integrity
            try:
                # Create a temporary file for testing
                temp_file_id = str(uuid.uuid4())
                temp_path = os.path.join(self.settings.TEMP_DIR, f"test_{temp_file_id}")
                
                with open(temp_path, 'wb') as f:
                    f.write(content[:min(1024*1024, len(content))])  # First 1MB for testing
                
                # Try to extract text (limited)
                test_result = await self._extract_by_type(content[:1024*100], temp_path, detected_type, file.filename)
                
                # Clean up test file
                try:
                    os.remove(temp_path)
                except:
                    pass
                
                if not test_result.get('text'):
                    return {
                        'valid': False,
                        'error': 'File appears to be corrupted or contains no extractable text'
                    }
                    
            except Exception as e:
                return {
                    'valid': False,
                    'error': f'File integrity check failed: {str(e)}'
                }
            
            return {
                'valid': True,
                'file_size': file_size,
                'detected_type': detected_type,
                'original_type': file.content_type
            }
            
        except Exception as e:
            logger.error("File validation failed", error=str(e))
            return {
                'valid': False,
                'error': f'Validation error: {str(e)}'
            }
    
    def _is_potentially_malicious(self, content: bytes, filename: str) -> bool:
        """Enhanced malicious content detection."""
        # Check for suspicious file extensions
        suspicious_extensions = ['.exe', '.bat', '.cmd', '.scr', '.vbs', '.jar', '.com', '.pif']
        if any(filename.lower().endswith(ext) for ext in suspicious_extensions):
            return True
        
        # Check for executable signatures
        if content.startswith(b'MZ'):  # PE executable
            return True
        
        if content.startswith(b'\x7fELF'):  # ELF executable
            return True
        
        # Check for script content in non-script files
        if b'<script' in content.lower() and not filename.lower().endswith(('.html', '.htm')):
            return True
        
        # Check for suspicious patterns
        suspicious_patterns = [
            b'powershell',
            b'cmd.exe',
            b'system32',
            b'<?php',
            b'#!/bin/',
            b'eval(',
            b'exec(',
        ]
        
        content_lower = content.lower()
        for pattern in suspicious_patterns:
            if pattern in content_lower and not filename.lower().endswith(('.txt', '.md', '.html', '.htm')):
                return True
        
        return False
    
    async def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get comprehensive file information."""
        try:
            path = Path(file_path)
            
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Basic file info
            stat = path.stat()
            
            # Read content for type detection
            async with aiofiles.open(file_path, 'rb') as f:
                content = await f.read(1024)  # First 1KB for type detection
            
            detected_type = self._detect_content_type(content, path.name)
            
            return {
                'path': str(path.absolute()),
                'name': path.name,
                'size': stat.st_size,
                'created': stat.st_ctime,
                'modified': stat.st_mtime,
                'detected_type': detected_type,
                'extension': path.suffix.lower(),
                'is_supported': detected_type in self.settings.SUPPORTED_UPLOAD_TYPES
            }
            
        except Exception as e:
            logger.error("Failed to get file info", file_path=file_path, error=str(e))
            raise
    
    async def cleanup_old_files(self, max_age_hours: int = 24):
        """Clean up old temporary files."""
        try:
            import time
            current_time = time.time()
            cutoff_time = current_time - (max_age_hours * 3600)
            
            temp_dir = Path(self.settings.TEMP_DIR)
            deleted_count = 0
            
            for file_path in temp_dir.iterdir():
                if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
                    try:
                        file_path.unlink()
                        deleted_count += 1
                    except Exception as e:
                        logger.warning("Failed to delete old file", 
                                     file=str(file_path), error=str(e))
            
            logger.info("Cleanup completed", deleted_files=deleted_count)
            
        except Exception as e:
            logger.error("Cleanup failed", error=str(e))
    
    def cleanup(self):
        """Clean up resources and temporary files."""
        try:
            for temp_file in self.temp_files.copy():
                try:
                    Path(temp_file).unlink(missing_ok=True)
                    self.temp_files.remove(temp_file)
                except Exception as e:
                    logger.warning("Failed to cleanup temp file", 
                                 file=temp_file, error=str(e))
            
            self.processing_cache.clear()
            logger.info("File handler cleanup completed")
            
        except Exception as e:
            logger.error("File handler cleanup failed", error=str(e))